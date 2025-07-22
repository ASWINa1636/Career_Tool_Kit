#---------------------------------THIS CODE WILL NOT STORE IN THE DATABASE-----------------------------------------------
from datetime import datetime
from docx import Document
from docx2pdf import convert
import fitz  # PyMuPDF
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os


def extract_text_from_file(file_path):
    if file_path.endswith('.pdf'):
        text = ""
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        return text
    elif file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return ""
        
# Score resume based on content
def score_resume(file_path):
    text = extract_text_from_file(file_path).lower()
    score = 0

    # --- Skills Section (Max 30) ---
    skill_keywords = ["Python", "Java", "JavaScript", "C++", "C#", "Ruby", "Go", "Swift", "Kotlin", "TypeScript", "PHP", "Rust",
    "Scala", "Perl", "Haskell", "Lua", "Linux", "Windows", "macOS", "UNIX", "MATLAB", "Power Systems", "HTML", 
    "CSS", "React", "Angular", "Vue.js", "Node.js", "Django", "Flask", "Spring Boot", "ASP.NET", "Laravel",
    "Machine Learning", "Deep Learning", "Data Science", "TensorFlow", "PyTorch", "Keras", "Pandas", "NumPy", 
    "Scikit-Learn", "R", "Matplotlib", "Seaborn", "OpenAI API", "Natural Language Processing", "Computer Vision",
    "Big Data", "SQL", "PostgreSQL", "MongoDB", "Firebase", "AWS", "Azure", "Google Cloud", "Docker", "Kubernetes",
    "Ethical Hacking", "Penetration Testing", "Cryptography", "Network Security", "SOC Analyst", "Malware Analysis",
    "Reverse Engineering", "CI/CD", "Jenkins", "Terraform", "Ansible", "Git", "GitHub Actions", "GitLab CI", 
    "Bash Scripting", "PowerShell", "Agile", "Scrum", "Kanban", "JIRA", "Trello", "Confluence", "PCB Design", 
    "Web Development", "Mobile Development", "Word", "Excel", "PowerPoint", "Outlook", "Tableau", "Power BI", 
    "Apache Spark", "Hadoop", "Kafka", "Elasticsearch", "GraphQL", "REST APIs", "SOAP", "Microservices", 
    "DevOps", "System Administration", "Virtualization", "VMware", "Hyper-V", "Cloud Security", "IoT", 
    "Embedded Systems", "Arduino", "Raspberry Pi", "Blockchain", "Solidity", "UI/UX Design", "Figma", "Adobe XD", 
    "Photoshop", "Illustrator", "Blender", "3D Modeling", "Game Development", "Unity", "Unreal Engine", "OpenGL", 
    "WebAssembly", "Quantum Computing", "Statistics", "Probability", "Linear Algebra", "Data Visualization", 
    "ETL Processes", "Data Warehousing", "Snowflake", "Redshift", "DynamoDB", "Cassandra", "Neo4j", "Redis", 
    "Load Balancing", "NGINX", "Apache", "Incident Response", "Forensic Analysis", "Threat Hunting", 
    "Cybersecurity Frameworks", "NIST", "ISO 27001", "GDPR Compliance", "Project Management", "PMP", "Lean Six Sigma", 
    "Technical Writing", "Public Speaking", "Team Leadership", "Conflict Resolution", "Time Management", 
    "Customer Relationship Management (CRM)", "Salesforce", "SAP", "ERP Systems", "Supply Chain Management", 
    "Digital Marketing", "SEO", "SEM", "Content Management Systems (CMS)", "WordPress", "Shopify", "Magento", 
    "Augmented Reality (AR)", "Virtual Reality (VR)", "Robotics", "ROS (Robot Operating System)", "PLC Programming", 
    "AutoCAD", "SolidWorks", "Finite Element Analysis (FEA)", "Computational Fluid Dynamics (CFD)", "Simulink", 
    "VLSI Design", "Verilog", "VHDL", "FPGA Programming", "Signal Processing", "Image Processing", "Audio Engineering", 
    "Penetration Testing Tools (Metasploit, Burp Suite)", "Wireshark", "Nmap", "Splunk", "SIEM", "Log Analysis", 
    "Chaos Engineering", "Site Reliability Engineering (SRE)", "Monitoring Tools (Prometheus, Grafana)", 
    "Version Control Systems", "Subversion (SVN)", "Mercurial", "Test Automation", "Selenium", "Cypress", 
    "Postman", "Unit Testing", "Integration Testing", "Performance Testing", "Load Testing", "Stress Testing", 
    "Behavior-Driven Development (BDD)", "Test-Driven Development (TDD)", "Pair Programming", "Code Review", 
    "Documentation", "API Design", "OAuth", "JWT", "Microfrontend", "Serverless Architecture", "Edge Computing", 
    "Bioinformatics", "Genomics", "Proteomics", "Molecular Modeling", "Chemoinformatics", "Financial Modeling", 
    "Risk Analysis", "Algorithm Design", "Data Structures", "Competitive Programming", "Parallel Computing", 
    "Distributed Systems", "Graph Theory", "Optimization", "Simulation", "Forecasting", "Econometrics", 
    "Geospatial Analysis", "GIS (Geographic Information Systems)", "Remote Sensing", "Satellite Imagery Analysis", 
    "Drone Technology", "Aeronautical Engineering", "Mechanical Design", "Thermodynamics", "Materials Science", 
    "Nanotechnology", "Renewable Energy Systems", "Solar Technology", "Wind Energy", "Battery Systems", 
    "Electrical Engineering", "Control Systems", "Power Electronics", "RF Engineering", "Antenna Design", 
    "Satellite Communications", "5G Technology", "Network Protocols", "TCP/IP", "DNS Management", "VPN Configuration", 
    "Customer Support", "Technical Support", "ITIL", "ServiceNow", "Help Desk Management", "Change Management", 
    "Disaster Recovery", "Business Continuity Planning", "Stakeholder Management", "Negotiation", "Critical Thinking", 
    "Problem Solving", "Emotional Intelligence", "Adaptability", "Cross-Functional Collaboration", "Mentoring", 
    "Training & Development", "Instructional Design", "E-Learning Development", "LMS (Learning Management Systems)"]
    skill_count = sum(1 for skill in skill_keywords if skill in text)
    score += min(skill_count * 3, 30)

    # --- Education (Max 20) ---
    if re.search(r'\b(b\.tech|btech|bachelor|be|mtech|m\.tech|msc|mca|bsc|degree)\b', text):
        score += 15
    if 'cgpa' in text or 'percentage' in text:
        score += 5

    # --- Experience (Max 20) ---
    if 'intern' in text or 'experience' in text or 'project' in text:
        score += 15
    if 'company' in text or 'organization' in text:
        score += 5

    # --- Contact Info (Max 15) ---
    if re.search(r'[\w\.-]+@[\w\.-]+', text):  # Email
        score += 8
    if re.search(r'\b\d{10}\b', text):  # Phone number
        score += 7

    # --- Structure (Max 15) ---
    if all(section in text for section in ['skills', 'education', 'experience']):
        score += 10
    elif any(section in text for section in ['skills', 'education', 'experience']):
        score += 5

    return min(score, 100)


# ---------------------- ATS Scanner ---------------------- #
def ats_scanner():
    while True:
        print("\n------------------------------------")
        print("         ATS SCANNER")
        print("------------------------------------")
        print("1. Paste Resume File Path and Scan")
        print("2. Back to Main Menu")

        sub_choice = input("Enter your choice: ")

        if sub_choice == '1':
            import string
            file_path = ''.join(ch for ch in input("Enter full path to resume (PDF/DOCX): ") if ch in string.printable).strip('"').strip()

            score = score_resume(file_path)
            print(f"\nScanning resume...\nResume Score: {score}/100 ✅")
            input("Press Enter to continue...")

        elif sub_choice == '2':
            break
        else:
            print("Invalid choice. Try again.")

            
# ---------------------- Resume Section ---------------------- #

def resume_creator():
    while True:
        print("\n------------------------------------")
        print("       RESUME MANAGER")
        print("------------------------------------")
        print("1. Create New Resume")
        print("2. Back to Main Menu")
        sub_choice = input("Enter your choice: ")
            
        if sub_choice == '1':
            name = input("Full Name: ")
            email = input("Email: ")
            phone = input("Phone Number: ")
            github = input("GitHub Link (or type 'skip'): ")
            linkedin = input("LinkedIn Link (or type 'skip'): ")
            portfolio = input("Portfolio Link (or type 'skip'): ")
    
            # Link - multiple entries
            custom = input("Custom Link (or type 'skip'): ")
            custom_label = ""
            if custom.lower() != 'skip':
                custom_label = input("Enter label for your Custom Link (e.g., Website): ").strip()

            # Create the Document   
            doc = Document()
            section = doc.sections[0]
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
    
            name_para = doc.add_paragraph(name)
            name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            name_para.runs[0].font.size = Pt(20)
            name_para.paragraph_format.space_after = Pt(6)
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            contact_para.paragraph_format.space_after = Pt(6)
            contact_para.add_run(f"{phone} | {email}")

            #Introducing hyperlink
            def add_hyperlink(paragraph, url, text):
                    part = paragraph.part
                    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
                    hyperlink = OxmlElement("w:hyperlink")
                    hyperlink.set(qn("r:id"), r_id)
                    new_run = OxmlElement("w:r")
                    rPr = OxmlElement("w:rPr")
                    color = OxmlElement("w:color")
                    color.set(qn("w:val"), "0000FF")
                    rPr.append(color)
                    underline = OxmlElement("w:u")
                    underline.set(qn("w:val"), "single")
                    rPr.append(underline)
                    new_run.append(rPr)
                    t = OxmlElement("w:t")
                    t.text = text
                    new_run.append(t)
                    hyperlink.append(new_run)
                    paragraph._p.append(hyperlink)

            if github.lower() != 'skip':
                contact_para.add_run(" | ")
                add_hyperlink(contact_para, github, "GitHub")
            if linkedin.lower() != 'skip':
                contact_para.add_run(" | ")
                add_hyperlink(contact_para, linkedin, "LinkedIn")
            if portfolio.lower() != 'skip':
                contact_para.add_run(" | ")
                add_hyperlink(contact_para, portfolio, "Portfolio")
            if custom.lower() != 'skip' and custom_label:
                contact_para.add_run(" | ")
                add_hyperlink(contact_para, custom, custom_label)  
        
            # About Section
            about = input("About You (or type 'skip'): ")
            
            if about.lower() != 'skip':
                doc.add_heading("About Me", level=1)
                doc.add_paragraph(about)              

            # Education - multiple entries
            education = []
            while True:
                edu = input("\nEducation (e.g., B.Tech, XYZ University, 20XX-20XX) or skip: ")
                if edu.lower() == 'skip':
                    break
                CGPA = input("Enter your CGPA or type 'skip': ")
                percent = input("Enter your percentage or 'skip': ")
                if CGPA.lower() == 'skip' and percent.lower() == 'skip':
                    education_str = f"{edu}"
                elif CGPA.lower() == 'skip':
                    education_str = f"{edu}\nPercentage: {percent}"
                elif percent.lower() == 'skip':
                    education_str = f"{edu}\nCGPA: {CGPA}"
                else:
                    education_str = f"{edu}\nCGPA: {CGPA}\nPercentage: {percent}"
                education.append(education_str)

            if education:
                doc.add_heading("Education", level=1)
                for e in education:
                    doc.add_paragraph(e, style='List Bullet')

            # Experience - multiple entries
            experiences = []
            while True:
                exp = input("\nEnter Experience (Role, Company, Duration) or 'skip': ")
                if exp.lower() == 'skip':
                    break
                experiences.append(exp)

            if experiences:
                doc.add_heading("Experience", level=1)
                for e in experiences:
                    doc.add_paragraph(e, style='List Bullet')

            # Projects - multiple entries
            projects = []
            while True:
                print("\n--- Add New Project (or type 'skip' to finish) ---")
                title = input("Project Title: ")
                if title.lower() == 'skip':
                    break
                desc = input("Description: ")
                tech = input("Tech Used or type 'skip': ")
                if tech.lower() == 'skip':
                    project_str = f"{title}\n {desc}"
                else:
                    project_str = f"{title}\n {desc}\n Tech Used: {tech}"
                projects.append(project_str)

            if projects:
                doc.add_heading("Projects", level=1)
                for p in projects:
                    doc.add_paragraph(p, style='List Bullet')                

            # Skills (structured)
            languages = input("Languages (e.g., Python, C++, SQL ) or (type 'skip'): ")
            tools = input("Developer Tools (e.g., Git, GitHub) or (type 'skip'): ")
            libraries = input("Libraries (e.g., Pandas, PyTorch) or (type 'skip'): ")
            soft_skills = input("Soft Skills (e.g., Communication, Teamwork) or (type 'skip'): ")

            # Custom skill sections - multiple entries
            custom_skills = []
            while True:
                add_custom = input("Do you want to add a custom skill section? (y/n): ").lower()
                if add_custom != 'y':
                    break
                title = input("  Section Title (e.g., Frameworks, Cloud Tools): ").strip()
                content = input("  List the skills/tools under this section: ").strip()
                custom_skills.append((title, content))    

            doc.add_heading("Skills", level=1)
            if languages.lower() != 'skip':
                doc.add_paragraph(f"•  Languages: {languages}")
            if tools.lower() != 'skip':
                doc.add_paragraph(f"•  Developer Tools: {tools}")
            if libraries.lower() != 'skip':
                doc.add_paragraph(f"•  Libraries: {libraries}")
            if soft_skills.lower() != 'skip':
                doc.add_paragraph(f"•  Soft Skills: {soft_skills}")
            for title, content in custom_skills:
                doc.add_paragraph(f"•  {title}: {content}")

            #certification - multiple entries
            certifications = []
            while True:
                cert = input("Enter Certifications or 'skip': ")
                if cert.lower() == 'skip':
                    break
                certifications.append(cert)

            if certifications:
               doc.add_heading("Certifications", level=1)
            for c in certifications:
                doc.add_paragraph(c, style='List Bullet') 

            filename = input("Enter filename (without extension): ").strip()
            
            # Create resume folder if it doesn't exist
            resume_folder = "D:/Carrer_toolKit/resume"
            os.makedirs(resume_folder, exist_ok=True)
            
            save_path = os.path.join(resume_folder, f"{filename}.docx")

            doc.save(save_path)
            convert(save_path)

            print(f"✅ Resume saved as {filename}.docx and {filename}.pdf in /resume/")
            input("Press Enter to continue...")
            continue

        elif sub_choice == '2':
            break
        else:
            print("Invalid choice. Try again.")

# ---------------------- SRM CGPE Calculator ---------------------- #

def cgpa_calculator():
    global conn, cursor
    print("\n------------------------------------")
    print("        SRM CGPA CALCULATOR")
    print("------------------------------------")

    
    total_semesters = int(input("Enter number of semesters completed: "))
    gpas = []
    for i in range(1, total_semesters + 1):
        gpa = float(input(f"Enter GPA for Semester {i}: "))
        gpas.append(gpa)
    if not gpas:
        print("❌ No GPA values entered.")
        return
    cgpa = sum(gpas) / len(gpas)
    print(f"\n🎓 Your CGPA after {total_semesters} semesters is: {cgpa:.2f}")
    input("Press Enter to continue...")

# ---------------------- Menu ---------------------- #
def print_menu():
    now = datetime.now()
    print("====================================")
    print("     SMART CAREER TOOLKIT")
    print("====================================")
    print(f"TIME: {now.strftime('%I:%M %p')}")
    print(f"DATE: {now.strftime('%d/%m/%Y')}")
    print("------------------------------------")
    print("1. ATS Scanner")
    print("2. Create Your Resume")
    print("3. SRM CGPA Calculator")
    print("5. Exit")

while True:
    print_menu()
    choice = input("\nEnter your choice: ")

    if choice == '1':
        print(">> [ATS Scanner Selected]")
        ats_scanner()  #  Call the scanner function here
    elif choice == '2':
        print(">> [Resume Creator Selected]")
        resume_creator()
    elif choice == '3':
        print(">> [CGPA Calculator Selected]")
        cgpa_calculator()
    elif choice == '5':
        print(">> Exiting... Goodbye!")
        break
    else:
        print("Invalid choice! Try again.")
